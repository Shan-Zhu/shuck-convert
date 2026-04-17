"""Markdown → DOCX conversion core logic.

Extracted from md2docx/md2docx.py, stripped of Windows UI (ctypes message boxes).
Requires pandoc to be installed and in PATH.
"""

import subprocess
import os
import re
import tempfile


# ── Table styling helpers ────────────────────────────────────────────────


def _make_border(val, sz="12", space="0", color="000000"):
    """创建一个边框 XML 元素。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    border = OxmlElement(f"w:{val}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), sz)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    return border


def _apply_three_line_table(table):
    """将表格设为三线表：顶线、表头底线（粗），底线（粗），其余无边框。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl

    tPr = tbl.tblPr
    if tPr is None:
        tPr = OxmlElement("w:tblPr")
        tbl.insert(0, tPr)

    for old_borders in tPr.findall(qn("w:tblBorders")):
        tPr.remove(old_borders)

    borders = OxmlElement("w:tblBorders")
    borders.append(_make_border("top", sz="12"))
    borders.append(_make_border("bottom", sz="12"))
    for side in ("left", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tPr.append(borders)

    if len(table.rows) > 1:
        for cell in table.rows[0].cells:
            tcPr = cell._element.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell._element.insert(0, tcPr)
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tc_borders = OxmlElement("w:tcBorders")
            tc_borders.append(_make_border("bottom", sz="12"))
            tcPr.append(tc_borders)


# ── Post-processing ─────────────────────────────────────────────────────


def apply_styles(docx_path):
    """后处理 DOCX：统一字体、行距、对齐、颜色，移除水平线。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement  # noqa: F811

    FONT_SIZE = Pt(12)

    doc = Document(docx_path)

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 设置默认段落格式
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 遍历所有段落
    paragraphs_to_remove = []
    for para in doc.paragraphs:
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                if not para.text.strip():
                    paragraphs_to_remove.append(para)
                    continue
                pPr.remove(pBdr)

        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = FONT_SIZE
            run.font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            rPr = r.find(qn("w:rPr"))
            if rPr is None:
                rPr = r.makeelement(qn("w:rPr"), {})
                r.insert(0, rPr)
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rPr.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")

    for para in paragraphs_to_remove:
        p = para._element
        p.getparent().remove(p)

    for table in doc.tables:
        _apply_three_line_table(table)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 2.0
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = FONT_SIZE
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        r = run._element
                        rPr = r.find(qn("w:rPr"))
                        if rPr is None:
                            rPr = r.makeelement(qn("w:rPr"), {})
                            r.insert(0, rPr)
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            rFonts = rPr.makeelement(qn("w:rFonts"), {})
                            rPr.insert(0, rFonts)
                        rFonts.set(qn("w:eastAsia"), "宋体")

    for s in doc.styles:
        if hasattr(s, "font") and s.font is not None:
            s.font.name = "Times New Roman"
            s.font.size = FONT_SIZE
            s.font.color.rgb = RGBColor(0, 0, 0)
            if hasattr(s, "element") and hasattr(s.element, "rPr") and s.element.rPr is not None:
                s.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if hasattr(s, "paragraph_format") and s.paragraph_format is not None:
            s.paragraph_format.line_spacing = 2.0
            s.paragraph_format.space_before = Pt(0)
            s.paragraph_format.space_after = Pt(0)

    doc.save(docx_path)


# ── Preprocessing ────────────────────────────────────────────────────────


def _find_image(filename, md_dir):
    """在 md 所在目录及其子目录中查找图片文件，返回相对路径。"""
    candidate = os.path.join(md_dir, filename)
    if os.path.isfile(candidate):
        return filename
    for root, _dirs, files in os.walk(md_dir):
        if filename in files:
            return os.path.relpath(os.path.join(root, filename), md_dir).replace('\\', '/')
    return filename


def _preprocess_obsidian_images(md_text, md_dir):
    """将 Obsidian 图片语法 ![[file|size]] 转换为标准 Markdown 图片语法。"""
    def _replace_img(m):
        content = m.group(1)
        parts = content.split('|', 1)
        filename = parts[0].strip()
        rel_path = _find_image(filename, md_dir)
        alt = os.path.splitext(filename)[0]
        return f'![{alt}]({rel_path})'

    return re.sub(r'!\[\[([^\]]+)\]\]', _replace_img, md_text)


def _preprocess_footnotes(md_text):
    """将各种引文/脚注语法转为行内文本，避免 pandoc 生成脚注。"""
    # 1) pandoc 行内脚注: ^[...] → [...]
    md_text = re.sub(r'\^\[([^\]]+)\]', r'[\1]', md_text)

    # 2) LaTeX 上标引用: ^{...} → [...]
    md_text = re.sub(r'\^\{([^}]+)\}', r'[\1]', md_text)

    # 3) 标准 markdown 脚注
    footnote_def_re = re.compile(
        r'^\[\^([^\]]+)\]:\s*(.*(?:\n(?![\[\n])(?:[ \t]+.*))*)',
        re.MULTILINE,
    )
    definitions = {}
    for m in footnote_def_re.finditer(md_text):
        fid = m.group(1)
        content = re.sub(r'\n[ \t]+', ' ', m.group(2)).strip()
        definitions[fid] = content

    if not definitions:
        return md_text

    cleaned = footnote_def_re.sub('', md_text)

    order = []
    def _replace_ref(m):
        fid = m.group(1)
        if fid not in order:
            order.append(fid)
        num = order.index(fid) + 1
        return f'[{num}]'

    cleaned = re.sub(r'\[\^([^\]]+)\]', _replace_ref, cleaned)
    cleaned = cleaned.rstrip('\n') + '\n'

    if order:
        ref_lines = ['\n---\n']
        for i, fid in enumerate(order, 1):
            content = definitions.get(fid, '')
            ref_lines.append(f'[{i}] {content}')
        cleaned += '\n'.join(ref_lines) + '\n'

    return cleaned


# ── Public API ───────────────────────────────────────────────────────────


def convert_markdown_to_docx(file_path: str) -> dict:
    """Convert a Markdown file to a styled DOCX document.

    Requires pandoc (https://pandoc.org/installing.html).

    Returns:
        dict with keys: output_path, message

    Raises:
        FileNotFoundError: if file doesn't exist
        RuntimeError: if pandoc is missing or conversion fails
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    docx_path = os.path.splitext(file_path)[0] + ".docx"
    md_dir = os.path.dirname(os.path.abspath(file_path))

    # Check pandoc
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
    except FileNotFoundError:
        raise RuntimeError(
            "pandoc not found. Install from https://pandoc.org/installing.html"
        )

    # Read and preprocess
    with open(file_path, "r", encoding="utf-8") as f:
        md_text = f.read()
    md_text = _preprocess_obsidian_images(md_text, md_dir)
    md_text = _preprocess_footnotes(md_text)

    # Write temp file and run pandoc
    tmp_md = None
    try:
        tmp_fd, tmp_md = tempfile.mkstemp(suffix=".md", dir=md_dir)
        with os.fdopen(tmp_fd, "w", encoding="utf-8") as f:
            f.write(md_text)

        result = subprocess.run(
            [
                "pandoc", tmp_md, "-o", docx_path,
                "--from", "markdown", "--to", "docx",
                "--resource-path", md_dir,
            ],
            capture_output=True,
            text=True,
            check=True,
            cwd=md_dir,
        )
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"pandoc conversion failed: {e.stderr}")
    finally:
        if tmp_md and os.path.exists(tmp_md):
            os.unlink(tmp_md)

    # Apply academic styles
    try:
        apply_styles(docx_path)
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    return {
        "output_path": docx_path,
        "message": f"Converted successfully: {docx_path}",
    }
